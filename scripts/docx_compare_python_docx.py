from __future__ import annotations

import argparse
import re
import sys
from collections import Counter
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[1]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="backslashreplace")
    sys.stderr.reconfigure(encoding="utf-8", errors="backslashreplace")
except Exception:
    pass


def _norm(s: str) -> str:
    s = s.replace("\u00a0", " ")
    s = re.sub(r"<<MT_[A-Za-z0-9_:\\-]{1,64}>>", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _extract_python_docx(docx_path: Path) -> list[str]:
    py_docx_src = REPO_ROOT / "python-docx-master" / "src"
    if str(py_docx_src) not in sys.path:
        sys.path.insert(0, str(py_docx_src))
    import docx  # type: ignore

    d = docx.Document(str(docx_path))
    out: list[str] = []

    def add_paragraphs(paragraphs) -> None:
        for p in paragraphs:
            t = _norm(p.text or "")
            if t:
                out.append(t)

    add_paragraphs(d.paragraphs)

    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                add_paragraphs(cell.paragraphs)

    for section in d.sections:
        add_paragraphs(section.header.paragraphs)
        add_paragraphs(section.footer.paragraphs)

    return out


def _extract_doctranslator_lxml(docx_path: Path) -> list[str]:
    from doctranslator.docx_package import parse_xml_part, read_docx
    from doctranslator.extract import extract_scopes_from_xml

    out: list[str] = []
    with read_docx(docx_path) as zin:
        xml_names = [
            info.filename
            for info in zin.infolist()
            if info.filename.lower().endswith(".xml") and info.file_size > 0
        ]
        xml_names.sort(
            key=lambda n: (0 if "word/document" in n.replace("\\", "/") else 1, n)
        )
        for name in xml_names:
            part = parse_xml_part(name, zin.read(name))
            scopes = extract_scopes_from_xml(name, part.tree)
            for sc in scopes:
                t = _norm((sc.surface_text or ""))
                if t:
                    out.append(t)
    return out


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("docx", type=Path)
    args = ap.parse_args()

    if not args.docx.exists():
        raise FileNotFoundError(args.docx)

    py = _extract_python_docx(args.docx)
    lx = _extract_doctranslator_lxml(args.docx)

    c_py = Counter(py)
    c_lx = Counter(lx)

    missing = list((c_py - c_lx).elements())
    extra = list((c_lx - c_py).elements())

    print(f"python-docx paragraphs: {len(py)} (unique={len(c_py)})")
    print(f"lxml scopes:          {len(lx)} (unique={len(c_lx)})")
    print(f"missing from lxml:    {len(missing)}")
    print(f"extra in lxml:        {len(extra)}")

    if missing:
        print("\n[missing examples]")
        for t in missing[:25]:
            print(f"- {t}")
    if extra:
        print("\n[extra examples]")
        for t in extra[:25]:
            print(f"- {t}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
