from __future__ import annotations

import argparse
import re
import sys
from collections import Counter
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[1]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="backslashreplace")
    sys.stderr.reconfigure(encoding="utf-8", errors="backslashreplace")
except Exception:
    pass


def _norm(s: str) -> str:
    s = s.replace("\u00a0", " ")
    s = re.sub(r"<<MT_[A-Za-z0-9_:\\-]{1,64}>>", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _extract_python_docx(docx_path: Path) -> list[str]:
    py_docx_src = REPO_ROOT / "python-docx-master" / "src"
    if str(py_docx_src) not in sys.path:
        sys.path.insert(0, str(py_docx_src))
    import docx  # type: ignore

    d = docx.Document(str(docx_path))
    out: list[str] = []

    def add_paragraphs(paragraphs) -> None:
        for p in paragraphs:
            t = _norm(p.text or "")
            if t:
                out.append(t)

    add_paragraphs(d.paragraphs)

    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                add_paragraphs(cell.paragraphs)

    for section in d.sections:
        add_paragraphs(section.header.paragraphs)
        add_paragraphs(section.footer.paragraphs)

    return out


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("a", type=Path)
    ap.add_argument("b", type=Path)
    args = ap.parse_args()

    if not args.a.exists():
        raise FileNotFoundError(args.a)
    if not args.b.exists():
        raise FileNotFoundError(args.b)

    a = _extract_python_docx(args.a)
    b = _extract_python_docx(args.b)

    c_a = Counter(a)
    c_b = Counter(b)

    missing = list((c_a - c_b).elements())
    extra = list((c_b - c_a).elements())

    print(f"A paragraphs: {len(a)} (unique={len(c_a)})")
    print(f"B paragraphs: {len(b)} (unique={len(c_b)})")
    print(f"missing from B: {len(missing)}")
    print(f"extra in B:     {len(extra)}")

    if missing:
        print("\n[missing examples]")
        for t in missing[:25]:
            print(f"- {t}")
    if extra:
        print("\n[extra examples]")
        for t in extra[:25]:
            print(f"- {t}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())

